from docx import Document
import argparse
import requests
import xmltodict


def setup_arguments():
    parser = argparse.ArgumentParser(
        description="Process committee meeting",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter
    )
    parser.add_argument(
        '-meeting',
        dest='meeting_id',
        type=int,
        required=True,
        help='Meeting ID used to fetch committee meeting data'
    )
    return parser.parse_args()


def get_meeting_data(meeting_id):
    url = f'https://legis.senado.leg.br/dadosabertos/reuniaocomissao/{meeting_id}'
    print(f'Fetching data from {url}')
    response = requests.get(url)
    # xml = response.content.decode('utf-8')
    # print(xml)
    return xmltodict.parse(response.text)


def generate_items_docx(meeting_data: dict, docx_name: str) -> None:
    document_name = f'{docx_name}.docx'
    document = Document()

    committee_nome = meeting_data["DetalheReuniao"]["reuniao"]["colegiadoCriador"]["nome"]
    committee_sigla = meeting_data["DetalheReuniao"]["reuniao"]["colegiadoCriador"]["sigla"]
    document.add_heading(f'{committee_nome} - {committee_sigla}', 0)

    document.add_paragraph(meeting_data["DetalheReuniao"]["reuniao"]["titulo"])
    document.add_paragraph(meeting_data["DetalheReuniao"]["reuniao"]["dataInicioFormatadaComObsHorario"])
    document.add_paragraph(meeting_data["DetalheReuniao"]["reuniao"]["local"])

    document.add_heading(meeting_data["DetalheReuniao"]["reuniao"]["partes"]["descricaoTipo"], 1)

    for item in meeting_data["DetalheReuniao"]["reuniao"]["partes"]["itens"]:
        document.add_heading(item['nomeFormatadoComOrdem'], level=2)
        document.add_paragraph(f'Ementa: {item["doma"]["ementa"]}')
        document.add_paragraph(f'Autoria: {item["doma"]["autoria"]}')
        # document.add_paragraph(f'Relatorio: {item["relatorio"]}')

        process_relatorias(document, item)
   
    document.save(document_name)
    print(f'Docx saved: {document_name}')


def process_relatorias(document, item):
    # Check if relatorias exist
    if not item.get('doma', {}).get('relatorias'):
        print(f"Item {item['nomeFormatadoComOrdem']} não possui relatoria")
        return

    relatorias = item['doma']['relatorias']
    # Ensure relatorias is always a list, even when there is only one relatoria
    relatorias_list = relatorias if isinstance(relatorias, list) else [relatorias]

    names = map(lambda r: r['relator']['parlamentar']['nomeComCargo'], relatorias_list)
    
    document.add_paragraph(f'Relatoria: {", ".join(names)}')


# ---------------------- MAIN ----------------------
print('---- Start')

args = setup_arguments()
meeting_data = get_meeting_data(args.meeting_id)
generate_items_docx(meeting_data, 'demo')

print('---- End')
