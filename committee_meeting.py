from docx import Document
import argparse
import requests
import xmltodict


def setup_arguments():
    parser = argparse.ArgumentParser(
        description="Process committee meeting",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter
    )
    parser.add_argument(
        '-meeting',
        dest='meeting_id',
        type=int,
        required=True,
        help='Meeting ID used to fetch committee meeting data'
    )
    return parser.parse_args()


def get_meeting_data(meeting_id):
    url = f'https://legis.senado.leg.br/dadosabertos/reuniaocomissao/{meeting_id}'
    print(f'Fetching data from {url}')
    response = requests.get(url)
    # xml = response.content.decode('utf-8')
    # print(xml)
    return xmltodict.parse(response.text)


def generate_items_docx(meeting_data: dict, docx_name: str) -> None:
    document_name = f'{docx_name}.docx'
    document = Document()

    document.add_heading(f'Reunião Comissão {meeting_data["DetalheReuniao"]["reuniao"]["codigo"]}', 0)
    document.add_paragraph(meeting_data["DetalheReuniao"]["reuniao"]["titulo"])

    document.add_heading('Partes', level=1)
    document.add_paragraph(meeting_data["DetalheReuniao"]["reuniao"]["partes"]["nome"])

    for item in meeting_data["DetalheReuniao"]["reuniao"]["partes"]["itens"]:
        document.add_heading(item['nomeFormatadoComOrdem'], level=2)
        document.add_paragraph(item['descricaoResultado'])

    document.save(document_name)
    print(f'Docx saved: {document_name}')


# ---------------------- MAIN ----------------------
print('---- Start')

args = setup_arguments()
meeting_data = get_meeting_data(args.meeting_id)
generate_items_docx(meeting_data, 'demo')

print('---- End')
